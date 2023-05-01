import streamlit as st
import docx2txt
from io import BytesIO
import base64
import docx
from audio_recorder_streamlit import audio_recorder
from google.cloud import speech_v1p1beta1 as speech
from google.cloud import texttospeech
import os
from pydub import AudioSegment
import streamlit.components.v1 as components




# Set up Streamlit and Google API
st.set_page_config(page_title="Question/Answer App", layout="wide")
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'key1.json'
lottie_file = "lottie.json"

# Transcribe audio using Google Speech-to-Text API
def transcribe_audio(mono_audio_bytes):
    client = speech.SpeechClient()
    audio = speech.RecognitionAudio(content=mono_audio_bytes)
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=48000,
        language_code="en-US",
        model="default"
    )

    response = client.recognize(config=config, audio=audio)
    transcript = ""

    for result in response.results:
        transcript += result.alternatives[0].transcript

    return transcript


def synthesize_speech(text):
    client = texttospeech.TextToSpeechClient()
    input_text = texttospeech.SynthesisInput(text=text)
    voice = texttospeech.VoiceSelectionParams(
        language_code="en-US", ssml_gender=texttospeech.SsmlVoiceGender.FEMALE
    )
    audio_config = texttospeech.AudioConfig(
        audio_encoding=texttospeech.AudioEncoding.LINEAR16
    )

    response = client.synthesize_speech(
        input=input_text, voice=voice, audio_config=audio_config
    )

    return response.audio_content


# Main application logic
def audio():
    button_style = """
        <style>
        section[data-testid="stSidebar"]  .stButton > button{
            display: inline-flex;
            -moz-box-align: center;
            align-items: center;
            -moz-box-pack: center;
            justify-content: center;
            font-weight: 400;
            padding: 0.25rem 0.75rem;
            border-radius: 0.25rem;
            margin: 0px;
            line-height: 1.6;
            color: inherit;
            width: auto;
            user-select: none;
            background-color: rgb(249, 249, 251);
            border: 1px solid rgba(49, 51, 63, 0.2);
            cursor: pointer;
            height: 35px;
         }
        section[data-testid="stSidebar"]  .stButton > button:hover { 
            border-color: rgb(255, 75, 75);
            color: rgb(255, 75, 75);
            }



        .stButton > button {
            color: white;
            background: #232121f5;
            width: 250px;
            height: 40px;
            font-size: 50px;
            }

            [data-testid="stVerticalBlock"] [data-testid="stHorizontalBlock"]:nth-child(1) .stButton {
            
                  position: fixed;
                    bottom: 60px;
              }   

        .stDownloadButton > button {
                    position: fixed;
                    bottom: 10px;
                    width: 250px;
                     
            }
            # [data-testid="stVerticalBlock"] div.css-budd4c.e1tzin5v0:nth-of-type(1) .stButton {
            #         position: fixed;
            #         bottom: 60px;
                    
                    
            } 
        </style>
        """
    st.markdown(button_style, unsafe_allow_html=True)


    # App title and description
    st.title("Question/Answer App")
    st.write("This app helps you to extract questions from a Word document, record your answers, and save them back into the document.")

    # Initialize session state for answers
    if "answers" not in st.session_state:
        st.session_state.answers = {}

    # Sidebar file uploader
    uploaded_file = st.file_uploader("Choose a Word document", type="docx")

    # Process uploaded file
    if uploaded_file is not None:
        # Extract text from the Word document
        file_contents = uploaded_file.read()
        text = docx2txt.process(BytesIO(file_contents))

        # Extract questions from the text
        questions = []
        for line in text.split("\n"):
            line = line.strip()
            if line.endswith("?"):
                questions.append(line)

        if not questions:
            st.warning(f":x: Sorry, I did not find any question in the document.")
        else:         

        # Display the current question
            question_index = st.number_input("Current Question:", min_value=1, max_value=len(questions), value=1, step=1,on_change=None,key="xy") - 1
            col1, col2 = st.columns([.7,1])
            col1.write(f"Total Number of Questions: {len(questions)}")
            
            
                
            if col1.button("Ask Next Question.🎧"):
                synthesized_audio = synthesize_speech(questions[question_index])

                audio_html = f"""
                <audio id="questionAudio" style="display:none;">
                    <source src="data:audio/wav;base64,{base64.b64encode(synthesized_audio).decode()}" type="audio/wav">
                </audio>
                <script>
                    document.querySelector("#questionAudio").play();
                </script>
                """

                components.html(audio_html, height=0)
            

                

            # Record audio answer
            recorder = audio_recorder(pause_threshold=60.0)
            audio_bytes = recorder

            # Process audio answer
            if audio_bytes:
                # Convert recorded audio to mono and the required sample rate
                stereo_audio = AudioSegment.from_wav(BytesIO(audio_bytes))
                mono_audio = stereo_audio.set_channels(1).set_frame_rate(48000)
                mono_audio_bytes = mono_audio.export(format="wav").read()

                # Convert audio to text button
                if st.button("Convert to text"):
                    answer = transcribe_audio(mono_audio_bytes)
                    col2.markdown("<b><span style='color:red; font-size: 20px;' >Answer:</span></b>",
                    unsafe_allow_html=True)
                    col2.write(answer)
                    
                    st.session_state.answers[question_index] = answer

            # Save answers to the document button
            with st.container():
                col1,col2=st.columns([.5,1])
                if col1.button("Save answers to the document"):
                    # Create a new Word document with the answers
                    doc = docx.Document(BytesIO(file_contents))
                    # ... (content remains the same)
                    title = doc.add_paragraph("Question and Answer")
                    title.runs[0].font.size = docx.shared.Pt(24)
                    title.runs[0].font.bold = True
                    doc.add_paragraph("")

                    for index, answer in st.session_state.answers.items():
                        question_paragraph = doc.add_paragraph()
                        question_paragraph.add_run("Question {}: ".format(index + 1)).bold = True
                        question_paragraph.add_run(questions[index])
                        answer_paragraph = doc.add_paragraph()
                        answer_paragraph.add_run("Answer {}: ".format(index + 1)).bold = True
                        answer_paragraph.add_run(answer)

                    try:
                        output_filename = "output_" + uploaded_file.name
                        with BytesIO() as output_file:
                            doc.save(output_file)
                            output_file.seek(0)
                            content = output_file.read()
                        st.success("Answers saved to the '{}'.".format(output_filename))

                        # Add the download button
                        col1.download_button(
                            label="Download updated document",
                            data=content,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    # Save the new document and display download button

                    except PermissionError:
                        st.error("Permission denied: Please close the current '{}' file before saving.".format(uploaded_file.name))

    else:
        st.error("Please upload a document before attempting to save answers.")


if __name__ == "__main__":
    audio()
